"""
Document Generator
==================
Main document generation engine.

Workflow:
1. Load DOCX template
2. Replace placeholders with actual data
3. Convert to PDF
4. Embed stamp and signature
5. Save and return path
"""

import os
import re
import uuid
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional, Any
from io import BytesIO

from .templates import TemplateManager, DocumentTemplate
from .stamp_service import StampService

# Try to import DOCX library
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("[DOCGEN] python-docx not available")


class DocumentGenerator:
    """
    Generates engineering documents from templates.

    Supports:
    - DOCX template processing
    - Placeholder replacement ({{key}} format)
    - PDF conversion via LibreOffice
    - Dynamic stamp embedding
    """

    OUTPUT_DIR = Path("outputs/documents")
    TEMP_DIR = Path("temp/documents")

    def __init__(self, templates_dir: str = "templates/documents"):
        """Initialize generator with template directory."""
        self.template_manager = TemplateManager(templates_dir)
        self.stamp_service = StampService()

        # Ensure output directories exist
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        self.TEMP_DIR.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        template_id: str,
        data: Dict[str, str],
        engineer_profile: Optional[Dict[str, Any]] = None,
        project_data: Optional[Dict[str, Any]] = None,
        output_format: str = "pdf",
        add_stamp: bool = True,
    ) -> Dict[str, Any]:
        """
        Generate document from template.

        Args:
            template_id: ID of template to use
            data: Dictionary of placeholder values
            engineer_profile: Engineer profile data (from API)
            project_data: Project data (from DB)
            output_format: 'pdf' or 'docx'
            add_stamp: Whether to add stamp/signature

        Returns:
            Dictionary with:
            - success: bool
            - path: Output file path
            - filename: Output filename
            - error: Error message if failed
        """
        try:
            # Get template
            template = self.template_manager.get_template(template_id)
            if not template:
                return {"success": False, "error": f"Template not found: {template_id}"}

            # Merge data sources
            merged_data = self._merge_data(data, engineer_profile, project_data)

            # Validate data
            missing = self.template_manager.validate_data(template_id, merged_data)
            if missing and "error" not in missing:
                # Allow generation with warnings for missing optional fields
                print(f"[DOCGEN] Warning: Missing fields: {missing}")

            # Get or create template file
            template_path = self._ensure_template_exists(template)
            if not template_path:
                return {"success": False, "error": "Template file not found"}

            # Generate unique output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = re.sub(r'[^\w\-]', '_', template.name_he)
            base_filename = f"{safe_name}_{timestamp}_{uuid.uuid4().hex[:8]}"

            # Process DOCX
            docx_path = self._process_docx(template_path, merged_data, base_filename)
            if not docx_path:
                return {"success": False, "error": "Failed to process DOCX template"}

            # Convert to PDF if requested
            if output_format == "pdf":
                pdf_path = self._convert_to_pdf(docx_path)
                if not pdf_path:
                    # Fallback: return DOCX if PDF conversion fails
                    print("[DOCGEN] PDF conversion failed, returning DOCX")
                    return {
                        "success": True,
                        "path": str(docx_path),
                        "filename": docx_path.name,
                        "format": "docx",
                        "warning": "PDF conversion failed",
                    }

                # Add stamp to PDF
                if add_stamp and engineer_profile:
                    pdf_path = self._add_stamp_to_pdf(pdf_path, engineer_profile)

                # Clean up temp DOCX
                docx_path.unlink(missing_ok=True)

                return {
                    "success": True,
                    "path": str(pdf_path),
                    "filename": pdf_path.name,
                    "format": "pdf",
                }
            else:
                return {
                    "success": True,
                    "path": str(docx_path),
                    "filename": docx_path.name,
                    "format": "docx",
                }

        except Exception as e:
            print(f"[DOCGEN] Error generating document: {e}")
            return {"success": False, "error": str(e)}

    def _merge_data(
        self,
        manual_data: Dict[str, str],
        engineer_profile: Optional[Dict[str, Any]],
        project_data: Optional[Dict[str, Any]],
    ) -> Dict[str, str]:
        """Merge all data sources into single placeholder dictionary."""
        merged = {}

        # Add system data
        now = datetime.now()
        merged["declaration_date"] = now.strftime("%d/%m/%Y")
        merged["declaration_date_hebrew"] = self._format_hebrew_date(now)
        merged["current_year"] = str(now.year)

        # Add engineer profile data
        if engineer_profile:
            merged["engineer_full_name"] = engineer_profile.get("full_name", "")
            merged["engineer_id"] = engineer_profile.get("id_number", "")
            merged["engineer_license"] = engineer_profile.get("engineer_license", "")
            merged["engineer_email"] = engineer_profile.get("email", "")
            merged["engineer_phone"] = engineer_profile.get("phone", "")

        # Add project data
        if project_data:
            merged["project_address"] = project_data.get("address", "")
            merged["gush_chalka"] = project_data.get("gush_chalka", "")
            merged["permit_number"] = project_data.get("permit_number", "")
            merged["project_name"] = project_data.get("name", "")
            merged["client_name"] = project_data.get("client_name", "")

        # Add manual data (overwrites if same keys)
        merged.update(manual_data)

        return merged

    def _format_hebrew_date(self, date: datetime) -> str:
        """Format date in Hebrew style."""
        months_he = [
            "ינואר", "פברואר", "מרץ", "אפריל", "מאי", "יוני",
            "יולי", "אוגוסט", "ספטמבר", "אוקטובר", "נובמבר", "דצמבר"
        ]
        return f"{date.day} ב{months_he[date.month - 1]} {date.year}"

    def _ensure_template_exists(self, template: DocumentTemplate) -> Optional[Path]:
        """Ensure template file exists, create sample if not."""
        path = self.template_manager.templates_dir / template.filename

        if not path.exists():
            print(f"[DOCGEN] Template not found, creating sample: {path}")
            if DOCX_SUPPORT:
                self.template_manager.create_sample_template(template.id)
                if path.exists():
                    return path

            # Create minimal text-based template as fallback
            self._create_minimal_template(path, template)

        return path if path.exists() else None

    def _create_minimal_template(self, path: Path, template: DocumentTemplate):
        """Create minimal template when DOCX is not available."""
        if DOCX_SUPPORT:
            doc = Document()
            doc.add_heading(template.name_he, level=1)
            doc.add_paragraph()

            for ph in template.placeholders:
                p = doc.add_paragraph()
                p.add_run(f"{ph.label_he}: ").bold = True
                p.add_run(f"{{{{{ph.key}}}}}")

            doc.save(str(path))
            print(f"[DOCGEN] Created minimal template: {path}")

    def _process_docx(
        self,
        template_path: Path,
        data: Dict[str, str],
        base_filename: str,
    ) -> Optional[Path]:
        """Process DOCX template and replace placeholders."""
        if not DOCX_SUPPORT:
            print("[DOCGEN] DOCX support not available")
            return None

        try:
            doc = Document(str(template_path))

            # Replace placeholders in all paragraphs
            for para in doc.paragraphs:
                self._replace_in_paragraph(para, data)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, data)

            # Replace in headers/footers
            for section in doc.sections:
                for para in section.header.paragraphs:
                    self._replace_in_paragraph(para, data)
                for para in section.footer.paragraphs:
                    self._replace_in_paragraph(para, data)

            # Save to temp directory
            output_path = self.TEMP_DIR / f"{base_filename}.docx"
            doc.save(str(output_path))

            print(f"[DOCGEN] Processed DOCX: {output_path}")
            return output_path

        except Exception as e:
            print(f"[DOCGEN] Error processing DOCX: {e}")
            return None

    def _replace_in_paragraph(self, para, data: Dict[str, str]):
        """Replace {{placeholders}} in a paragraph while preserving formatting."""
        # Check if paragraph contains any placeholder
        full_text = para.text
        if "{{" not in full_text:
            return

        # Replace all placeholders
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))

        # Update paragraph text (simple method - may lose some formatting)
        if para.runs:
            # Clear existing runs except first
            for run in para.runs[1:]:
                run.clear()
            # Set text in first run
            para.runs[0].text = full_text
        else:
            para.text = full_text

    def _convert_to_pdf(self, docx_path: Path) -> Optional[Path]:
        """Convert DOCX to PDF using LibreOffice."""
        try:
            output_dir = self.OUTPUT_DIR

            # Try LibreOffice conversion
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(output_dir),
                    str(docx_path),
                ],
                capture_output=True,
                timeout=60,
            )

            if result.returncode == 0:
                pdf_path = output_dir / f"{docx_path.stem}.pdf"
                if pdf_path.exists():
                    print(f"[DOCGEN] Converted to PDF: {pdf_path}")
                    return pdf_path

            print(f"[DOCGEN] LibreOffice conversion failed: {result.stderr.decode()}")
            return None

        except FileNotFoundError:
            print("[DOCGEN] LibreOffice not installed")
            return None
        except subprocess.TimeoutExpired:
            print("[DOCGEN] PDF conversion timed out")
            return None
        except Exception as e:
            print(f"[DOCGEN] PDF conversion error: {e}")
            return None

    def _add_stamp_to_pdf(
        self,
        pdf_path: Path,
        engineer_profile: Dict[str, Any],
    ) -> Path:
        """Add stamp and signature to PDF."""
        stamp_path = engineer_profile.get("stamp_signature_path")
        if not stamp_path or not os.path.exists(stamp_path):
            print("[DOCGEN] No stamp image available, skipping")
            return pdf_path

        try:
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()

            stamped_bytes = self.stamp_service.embed_stamp(
                pdf_bytes=pdf_bytes,
                stamp_image_path=stamp_path,
                engineer_name=engineer_profile.get("full_name", ""),
                engineer_id=engineer_profile.get("id_number", ""),
            )

            # Save stamped PDF
            stamped_path = pdf_path.with_name(f"{pdf_path.stem}_signed.pdf")
            with open(stamped_path, 'wb') as f:
                f.write(stamped_bytes)

            # Remove unsigned PDF
            pdf_path.unlink(missing_ok=True)

            print(f"[DOCGEN] Added stamp: {stamped_path}")
            return stamped_path

        except Exception as e:
            print(f"[DOCGEN] Error adding stamp: {e}")
            return pdf_path
