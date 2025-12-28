"""
Template Manager
================
Manages document templates with placeholder support.
"""

import os
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, field
from enum import Enum


class TemplateType(str, Enum):
    """Types of engineering document templates."""
    PLUMBING_AFFIDAVIT_AFTER = "plumbing_affidavit_after_execution"
    PLUMBING_COMPLETION = "plumbing_completion_certificate"
    ELECTRICAL_AFFIDAVIT = "electrical_affidavit"
    STRUCTURE_AFFIDAVIT = "structure_affidavit"
    FIRE_SAFETY = "fire_safety_approval"
    FORM_4 = "form_4"
    GENERAL = "general_declaration"


@dataclass
class TemplatePlaceholder:
    """Definition of a template placeholder."""
    key: str  # {{key}} format
    label_he: str
    label_en: str
    source: str  # 'engineer', 'project', 'system', 'manual'
    required: bool = True
    default: Optional[str] = None


@dataclass
class DocumentTemplate:
    """Document template definition."""
    id: str
    type: TemplateType
    name_he: str
    name_en: str
    filename: str
    placeholders: List[TemplatePlaceholder] = field(default_factory=list)
    description_he: str = ""
    description_en: str = ""
    category: str = "plumbing"


# =============================================================================
# TEMPLATE REGISTRY
# =============================================================================

# Standard placeholders used across templates
ENGINEER_PLACEHOLDERS = [
    TemplatePlaceholder("engineer_full_name", "שם מלא", "Full Name", "engineer"),
    TemplatePlaceholder("engineer_id", "תעודת זהות", "ID Number", "engineer"),
    TemplatePlaceholder("engineer_license", "מספר רישיון", "License Number", "engineer", required=False),
    TemplatePlaceholder("engineer_email", "אימייל", "Email", "engineer"),
    TemplatePlaceholder("engineer_phone", "טלפון", "Phone", "engineer"),
]

PROJECT_PLACEHOLDERS = [
    TemplatePlaceholder("project_address", "כתובת הפרויקט", "Project Address", "project"),
    TemplatePlaceholder("gush_chalka", "גוש/חלקה", "Block/Parcel", "project"),
    TemplatePlaceholder("permit_number", "מספר היתר", "Permit Number", "project", required=False),
    TemplatePlaceholder("project_name", "שם הפרויקט", "Project Name", "project", required=False),
    TemplatePlaceholder("client_name", "שם הלקוח", "Client Name", "project", required=False),
]

SYSTEM_PLACEHOLDERS = [
    TemplatePlaceholder("declaration_date", "תאריך", "Date", "system"),
    TemplatePlaceholder("declaration_date_hebrew", "תאריך עברי", "Hebrew Date", "system"),
    TemplatePlaceholder("current_year", "שנה", "Year", "system"),
    TemplatePlaceholder("signature_stamp", "חותמת וחתימה", "Stamp & Signature", "system"),
]

# Template definitions
TEMPLATES: Dict[str, DocumentTemplate] = {
    "plumbing_affidavit_after": DocumentTemplate(
        id="plumbing_affidavit_after",
        type=TemplateType.PLUMBING_AFFIDAVIT_AFTER,
        name_he="תצהיר מהנדס אינסטלציה לאחר ביצוע",
        name_en="Plumbing Engineer Affidavit - Post Construction",
        filename="plumbing_affidavit_after_execution.docx",
        placeholders=[
            *ENGINEER_PLACEHOLDERS,
            *PROJECT_PLACEHOLDERS,
            *SYSTEM_PLACEHOLDERS,
            TemplatePlaceholder("work_description", "תיאור העבודה", "Work Description", "manual"),
            TemplatePlaceholder("inspection_date", "תאריך בדיקה", "Inspection Date", "manual", required=False),
        ],
        description_he="תצהיר מהנדס על השלמת עבודות אינסטלציה בהתאם לתוכניות ולתקנים",
        description_en="Engineer declaration of plumbing work completion per plans and standards",
        category="plumbing",
    ),
    "plumbing_completion": DocumentTemplate(
        id="plumbing_completion",
        type=TemplateType.PLUMBING_COMPLETION,
        name_he="אישור גמר אינסטלציה",
        name_en="Plumbing Completion Certificate",
        filename="plumbing_completion_certificate.docx",
        placeholders=[
            *ENGINEER_PLACEHOLDERS,
            *PROJECT_PLACEHOLDERS,
            *SYSTEM_PLACEHOLDERS,
        ],
        description_he="אישור גמר עבודות אינסטלציה",
        description_en="Certificate of plumbing work completion",
        category="plumbing",
    ),
}


class TemplateManager:
    """Manages document templates and their placeholders."""

    def __init__(self, templates_dir: str = "templates/documents"):
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def get_template(self, template_id: str) -> Optional[DocumentTemplate]:
        """Get a template by ID."""
        return TEMPLATES.get(template_id)

    def list_templates(self, category: Optional[str] = None) -> List[DocumentTemplate]:
        """List all templates, optionally filtered by category."""
        templates = list(TEMPLATES.values())
        if category:
            templates = [t for t in templates if t.category == category]
        return templates

    def get_template_path(self, template_id: str) -> Optional[Path]:
        """Get the file path for a template."""
        template = self.get_template(template_id)
        if not template:
            return None
        path = self.templates_dir / template.filename
        return path if path.exists() else None

    def get_placeholders(self, template_id: str) -> List[TemplatePlaceholder]:
        """Get all placeholders for a template."""
        template = self.get_template(template_id)
        if not template:
            return []
        return template.placeholders

    def validate_data(self, template_id: str, data: Dict[str, str]) -> Dict[str, str]:
        """
        Validate provided data against template placeholders.
        Returns dict of missing required fields.
        """
        template = self.get_template(template_id)
        if not template:
            return {"error": "Template not found"}

        missing = {}
        for placeholder in template.placeholders:
            if placeholder.required and placeholder.key not in data:
                missing[placeholder.key] = f"Missing: {placeholder.label_he}"

        return missing

    def create_sample_template(self, template_id: str) -> bool:
        """
        Create a sample DOCX template with placeholders.
        Useful for initial setup.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            template = self.get_template(template_id)
            if not template:
                return False

            doc = Document()

            # Title
            title = doc.add_heading(template.name_he, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # Add placeholders
            for placeholder in template.placeholders:
                p = doc.add_paragraph()
                p.add_run(f"{placeholder.label_he}: ").bold = True
                p.add_run(f"{{{{{placeholder.key}}}}}")

            # Signature area
            doc.add_paragraph()
            doc.add_paragraph()
            sig = doc.add_paragraph()
            sig.add_run("חותמת וחתימה: ").bold = True
            sig.add_run("{{signature_stamp}}")

            # Save
            filepath = self.templates_dir / template.filename
            doc.save(str(filepath))
            print(f"[TEMPLATE] Created sample: {filepath}")
            return True

        except ImportError:
            print("[TEMPLATE] python-docx not installed, skipping sample creation")
            return False
        except Exception as e:
            print(f"[TEMPLATE] Error creating sample: {e}")
            return False
